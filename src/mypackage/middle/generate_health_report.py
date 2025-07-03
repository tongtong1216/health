from docx import Document
from docx.oxml.ns import qn
from datetime import datetime
from middle.report import Health_report
import os

def _set_document_font(document, font_name='宋体'):
    """设置整个文档默认字体为宋体"""
    # 设置文档主题字体
    try:
        theme_fonts = document.part.theme_part.theme.themeElements.fontScheme
        theme_fonts.minorFont.latin.typeface = font_name
        theme_fonts.majorFont.latin.typeface = font_name
    except Exception:
        pass
    
    # 设置默认段落字体
    styles = document.styles
    for style in styles:
        try:
            if style.type == 1:  # 段落样式
                style.font.name = font_name
                style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except AttributeError:
            continue
    
    # 设置主题字体
    settings = document.part.element.xpath('//w:themeFontLang')
    for setting in settings:
        setting.set(qn('w:val'), 'zh-CN')
        setting.set(qn('w:eastAsia'), font_name)
    
    # 确保所有运行都使用宋体
    for section in document.sections:
        # 页眉
        for header in [section.header]:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # 页脚
        for footer in [section.footer]:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def generate_health_report(username: str) -> None:
    """
    生成用户健康报告
    :param username: 用户名
    """
    user_data = Health_report.health_report(username)
    user_data['日期'] = datetime.now().strftime("%Y年%m月%d日")
    # 用户数据
    # user_data = {
    #     '标题': f"{username}健康报告",
    #     '个人信息': f"年龄:{age},性别:{gender},身高:{height}米,体重:{weight}千克",
    #     'bmi评估': f"BMI值{bmi_round},{bmi_eval}",
    #     '运动量评估': f"最近一周锻炼时间:{duration},运动量评估:{exercise_eval}",
    #     '睡眠情况评估': f"睡眠时间:{sleep_time},睡眠情况评估:{sleep_eval}",
    #     '静息心率状况评估': f"静息心率:{heart_rate},心率状况评估:{heart_eval}",
    #     '血压状况评估': f"血压值（收缩压/舒张压）:{blood_pressure_data},血压状况评估:{bp_eval}",
    #     '血糖状况评估': f"空腹血糖值:{blood_glucose},血糖状况评估:{glucose_eval}",
    #     '日期': datetime.now().strftime("%Y年%m月%d日")
    # }


    # 获取当前脚本所在的目录
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 自定义模板路径和输出目录（相对于脚本目录）
    template_path = os.path.join(script_dir, 'health_report_template.docx')
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'reports')

    # 加载模板
    doc =   Document(template_path)

    # 替换占位符
    for paragraph in doc.paragraphs:
        for key, value in user_data.items():
            placeholder = f"{{{{ {key} }}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # 替换页脚
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            for key, value in user_data.items():
                placeholder = f"{{{{ {key} }}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

    
    _set_document_font(doc, font_name='宋体')
    
    # 保存最终报告
    output_filename = os.path.join(output_dir, f"{username}_健康报告_{datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(output_filename)
    print(f"健康报告已生成: {output_filename}")